#!/usr/bin/env python3
"""
Extractor completo de documento DOCX para análisis jurídico PPMM-GEI
Enhanced Universal Framework v3.0 aplicado
"""

from docx import Document
import sys
from pathlib import Path

def extract_docx_content(file_path):
    """Extrae todo el contenido del DOCX incluyendo texto, tablas y estructura"""
    try:
        doc = Document(file_path)
        
        content = {
            'paragraphs': [],
            'tables': [],
            'full_text': [],
            'headings': [],
            'metadata': {}
        }
        
        # Extraer párrafos con estilo
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                para_info = {
                    'index': i,
                    'text': text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
                }
                content['paragraphs'].append(para_info)
                content['full_text'].append(text)
                
                if para_info['is_heading'] or len(text) < 200:
                    content['headings'].append(text)
        
        # Extraer tablas
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'index': i,
                'data': table_data
            })
        
        # Metadata
        try:
            props = doc.core_properties
            content['metadata'] = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else ''
            }
        except:
            content['metadata'] = {}
        
        return content
        
    except Exception as e:
        print(f"Error extrayendo DOCX: {str(e)}")
        return None

def analyze_legal_structure(content):
    """Analiza la estructura jurídica del documento"""
    analysis = {
        'articles': [],
        'definitions': [],
        'procedures': [],
        'institutions': [],
        'key_concepts': []
    }
    
    for para in content['paragraphs']:
        text = para['text']
        text_lower = text.lower()
        
        # Detectar artículos
        if any(marker in text_lower for marker in ['artículo', 'art.', 'art ']):
            analysis['articles'].append(text)
        
        # Detectar definiciones
        if any(marker in text_lower for marker in ['definición', 'definir', 'entiende por', 'se considera']):
            analysis['definitions'].append(text)
        
        # Detectar procedimientos
        if any(marker in text_lower for marker in ['procedimiento', 'trámite', 'solicitud', 'registro']):
            analysis['procedures'].append(text)
        
        # Detectar instituciones
        if any(marker in text_lower for marker in ['autoridad', 'ministerio', 'organismo', 'secretaría']):
            analysis['institutions'].append(text)
        
        # Conceptos clave carbono
        if any(marker in text_lower for marker in ['carbono', 'emisión', 'mitigación', 'gei', 'invernadero']):
            analysis['key_concepts'].append(text)
    
    return analysis

def main():
    docx_file = Path("/home/user/webapp/Analisis_Juridico_Proyecto_PPMM_GEI.docx")
    
    if not docx_file.exists():
        print(f"Error: No se encontró el archivo {docx_file}")
        return
    
    print("🔍 EXTRAYENDO CONTENIDO COMPLETO DEL DOCUMENTO DOCX...")
    content = extract_docx_content(docx_file)
    
    if not content:
        print("❌ Error al extraer contenido")
        return
    
    print(f"✅ Documento extraído exitosamente:")
    print(f"   - {len(content['paragraphs'])} párrafos")
    print(f"   - {len(content['tables'])} tablas")
    print(f"   - {len(content['headings'])} encabezados")
    
    # Análisis jurídico
    print("\n🏛️ ANALIZANDO ESTRUCTURA JURÍDICA...")
    legal_analysis = analyze_legal_structure(content)
    
    print(f"   - {len(legal_analysis['articles'])} artículos detectados")
    print(f"   - {len(legal_analysis['definitions'])} definiciones detectadas")
    print(f"   - {len(legal_analysis['procedures'])} procedimientos detectados")
    print(f"   - {len(legal_analysis['institutions'])} instituciones detectadas")
    print(f"   - {len(legal_analysis['key_concepts'])} conceptos clave detectados")
    
    # Guardar contenido completo
    output_file = Path("/home/user/webapp/documento_ppmm_gei_completo_extraido.txt")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("# DOCUMENTO PPMM-GEI EXTRAÍDO COMPLETO\n")
        f.write("# Enhanced Universal Framework v3.0 - Extracción Jurídica Completa\n\n")
        
        f.write("## METADATA\n")
        for key, value in content['metadata'].items():
            f.write(f"{key}: {value}\n")
        f.write("\n")
        
        f.write("## CONTENIDO COMPLETO\n\n")
        for para in content['paragraphs']:
            f.write(f"{para['text']}\n\n")
        
        if content['tables']:
            f.write("## TABLAS\n\n")
            for i, table in enumerate(content['tables']):
                f.write(f"### Tabla {i+1}\n")
                for row in table['data']:
                    f.write(" | ".join(row) + "\n")
                f.write("\n")
        
        f.write("## ANÁLISIS JURÍDICO AUTOMÁTICO\n\n")
        
        f.write("### ARTÍCULOS DETECTADOS\n")
        for article in legal_analysis['articles']:
            f.write(f"- {article}\n")
        f.write("\n")
        
        f.write("### DEFINICIONES DETECTADAS\n")
        for definition in legal_analysis['definitions']:
            f.write(f"- {definition}\n")
        f.write("\n")
        
        f.write("### PROCEDIMIENTOS DETECTADOS\n")
        for procedure in legal_analysis['procedures']:
            f.write(f"- {procedure}\n")
        f.write("\n")
        
        f.write("### INSTITUCIONES DETECTADAS\n")
        for institution in legal_analysis['institutions']:
            f.write(f"- {institution}\n")
        f.write("\n")
        
        f.write("### CONCEPTOS CLAVE CARBONO\n")
        for concept in legal_analysis['key_concepts']:
            f.write(f"- {concept}\n")
    
    print(f"\n💾 Contenido completo guardado en: {output_file}")
    
    # Mostrar primeros párrafos para verificación
    print("\n📋 PRIMEROS PÁRRAFOS EXTRAÍDOS:")
    for i, para in enumerate(content['paragraphs'][:5]):
        print(f"\n{i+1}. [{para['style']}] {para['text'][:200]}{'...' if len(para['text']) > 200 else ''}")

if __name__ == "__main__":
    main()